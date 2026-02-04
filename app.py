# Edited By Song Zichen, Quan Hai Middle School
from flask import Flask, render_template, request, jsonify, send_file, Response, stream_with_context
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import random
import json
import os
import re
from datetime import datetime, timedelta
import ollama
import threading
import platform
import subprocess
import uuid
from werkzeug.utils import secure_filename

# 简单的内存缓存实现
cache = {}
cache_expiry = {}
CACHE_DURATION = timedelta(minutes=5)

def get_cached_data(key):
    """获取缓存数据"""
    if key in cache and datetime.now() < cache_expiry.get(key, datetime.min):
        return cache[key]
    return None

def set_cached_data(key, data):
    """设置缓存数据"""
    cache[key] = data
    cache_expiry[key] = datetime.now() + CACHE_DURATION

def clear_cache():
    """清除所有缓存"""
    cache.clear()
    cache_expiry.clear()

app = Flask(__name__)

# 错题本配置
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp', 'pdf', 'doc', 'docx', 'txt', 'md', 'mp4', 'avi', 'mov', 'mkv', 'webm', 'mp3', 'wav'}
ERROR_NOTEBOOK_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'error_notebook.json')
MAX_CONTENT_LENGTH = 500 * 1024 * 1024  # 500MB 最大上传大小

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# 确保上传目录存在
os.makedirs(os.path.join(UPLOAD_FOLDER, 'questions'), exist_ok=True)
os.makedirs(os.path.join(UPLOAD_FOLDER, 'answers'), exist_ok=True)
os.makedirs(os.path.join(UPLOAD_FOLDER, 'videos'), exist_ok=True)

def allowed_file(filename):
    """检查文件扩展名是否允许"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_error_notebook_data():
    """获取错题本数据"""
    if os.path.exists(ERROR_NOTEBOOK_FILE):
        with open(ERROR_NOTEBOOK_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {'errors': []}

def save_error_notebook_data(data):
    """保存错题本数据"""
    with open(ERROR_NOTEBOOK_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def get_file_type(filename):
    """根据文件名获取文件类型"""
    if not filename:
        return 'unknown'
    ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    if ext in {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp'}:
        return 'image'
    elif ext == 'pdf':
        return 'pdf'
    elif ext in {'doc', 'docx'}:
        return 'word'
    elif ext in {'txt', 'md'}:
        return 'text'
    elif ext in {'mp4', 'avi', 'mov', 'mkv', 'webm'}:
        return 'video'
    elif ext in {'mp3', 'wav'}:
        return 'audio'
    return 'file'

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def generate_math_questions(config):
    data_file = os.path.join('data', 'math_config.json')
    if not os.path.exists(data_file):
        return []
    
    with open(data_file, 'r', encoding='utf-8') as f:
        math_config = json.load(f)
    
    grade_id = config.get('grade_id')
    category_id = config.get('category_id')
    type_id = config.get('type_id')
    count = config.get('count', 20)
    
    questions = []
    
    for grade in math_config.get('grade_levels', []):
        if grade.get('id') == grade_id:
            for category in grade.get('categories', []):
                if category.get('id') == category_id:
                    for type_info in category.get('types', []):
                        if type_info.get('id') == type_id:
                            for _ in range(count):
                                question = generate_single_math_question(grade_id, category_id, type_id)
                                questions.append(question)
                            break
                    break
            break
    
    return questions

def generate_single_math_question(grade_id, category_id, type_id):
    if grade_id == '1-2':
        return generate_grade_1_2_question(category_id, type_id)
    elif grade_id == '3-4':
        return generate_grade_3_4_question(category_id, type_id)
    elif grade_id == '5-6':
        return generate_grade_5_6_question(category_id, type_id)
    return {'question': '', 'answer': ''}

def generate_grade_1_2_question(category_id, type_id):
    if category_id == 'addition_within_20':
        if type_id == 'oral_calculation':
            a = random.randint(1, 18)
            b = random.randint(1, 20 - a)
            return {'question': f"{a} + {b} = ", 'answer': a + b}
        elif type_id == 'continuous_addition':
            a = random.randint(1, 10)
            b = random.randint(1, 10)
            c = random.randint(1, 10)
            return {'question': f"{a} + {b} + {c} = ", 'answer': a + b + c}
        elif type_id == 'continuous_subtraction':
            a = random.randint(10, 20)
            b = random.randint(1, 8)
            c = random.randint(1, a - b)
            return {'question': f"{a} - {b} - {c} = ", 'answer': a - b - c}
        elif type_id == 'mixed_addition_subtraction':
            a = random.randint(5, 15)
            b = random.randint(1, 8)
            c = random.randint(1, 6)
            return {'question': f"{a} + {b} - {c} = ", 'answer': a + b - c}
    elif category_id == 'addition_within_100':
        if type_id == 'oral_calculation_100':
            if random.choice([True, False]):
                a = random.randint(1, 9) * 10
                b = random.randint(1, 9) * 10
                return {'question': f"{a} + {b} = ", 'answer': a + b}
            else:
                a = random.randint(10, 99)
                b = random.randint(1, 9)
                op = random.choice(['+', '-'])
                if op == '+':
                    return {'question': f"{a} + {b} = ", 'answer': a + b}
                else:
                    return {'question': f"{a} - {b} = ", 'answer': a - b}
        elif type_id == 'written_calculation_100':
            if random.choice([True, False]):
                a = random.randint(10, 89)
                b = random.randint(1, 99 - a)
                return {'question': f"{a} + {b} = ", 'answer': a + b}
            else:
                a = random.randint(20, 99)
                b = random.randint(1, a)
                return {'question': f"{a} - {b} = ", 'answer': a - b}
    elif category_id == 'multiplication_division':
        if type_id == 'multiplication_table':
            a = random.randint(2, 9)
            b = random.randint(2, 9)
            return {'question': f"{a} × {b} = ", 'answer': a * b}
        elif type_id == 'division_by_table':
            b = random.randint(2, 9)
            answer = random.randint(2, 9)
            a = b * answer
            return {'question': f"{a} ÷ {b} = ", 'answer': answer}
    elif category_id == 'simple_mixed':
        if type_id == 'same_level':
            if random.choice([True, False]):
                a = random.randint(5, 20)
                b = random.randint(1, 10)
                c = random.randint(1, 10)
                return {'question': f"{a} + {b} - {c} = ", 'answer': a + b - c}
            else:
                a = random.randint(6, 36)
                b = random.randint(2, 6)
                c = random.randint(2, 6)
                return {'question': f"{a} ÷ {b} × {c} = ", 'answer': (a // b) * c}
    
    return {'question': '', 'answer': ''}

def generate_grade_3_4_question(category_id, type_id):
    if category_id == 'multi_digit_multiplication':
        if type_id == 'oral_multi':
            a = random.randint(1, 9) * random.choice([10, 100])
            b = random.randint(2, 9)
            return {'question': f"{a} × {b} = ", 'answer': a * b}
        elif type_id == 'written_multi':
            a = random.randint(100, 999)
            b = random.randint(2, 9)
            return {'question': f"{a} × {b} = ", 'answer': a * b}
    elif category_id == 'remainder_division':
        if type_id == 'remainder_calc':
            b = random.randint(3, 9)
            answer = random.randint(10, 20)
            a = b * answer + random.randint(1, b - 1)
            return {'question': f"{a} ÷ {b} = ", 'answer': f"{answer}……{a % b}"}
    elif category_id == 'two_digit_multiplication':
        if type_id == 'estimate_multi':
            a = random.randint(10, 99)
            b = random.randint(10, 99)
            return {'question': f"{a} × {b} ≈ ", 'answer': a * b}
        elif type_id == 'written_multi_2digit':
            a = random.randint(10, 99)
            b = random.randint(10, 99)
            return {'question': f"{a} × {b} = ", 'answer': a * b}
    elif category_id == 'one_digit_division':
        if type_id == 'zero_division':
            b = random.randint(3, 9)
            answer = random.choice([100, 200, 300, 400, 500, 600])
            a = answer * b
            return {'question': f"{a} ÷ {b} = ", 'answer': answer}
    elif category_id == 'four_operations':
        if type_id == 'with_parentheses':
            if random.choice([True, False]):
                a = random.randint(10, 50)
                b = random.randint(5, 30)
                c = random.randint(2, 9)
                return {'question': f"({a} + {b}) × {c} = ", 'answer': (a + b) * c}
            else:
                a = random.randint(20, 100)
                b = random.randint(6, 18)
                c = random.randint(2, 9)
                return {'question': f"{a} ÷ ({b} ÷ {c}) = ", 'answer': a // (b // c)}
    
    return {'question': '', 'answer': ''}

def generate_grade_5_6_question(category_id, type_id):
    if category_id == 'decimal_operations':
        if type_id == 'decimal_addition':
            if random.choice([True, False]):
                a = round(random.uniform(1, 10), 2)
                b = round(random.uniform(0.1, 5), 2)
                return {'question': f"{a} + {b} = ", 'answer': round(a + b, 2)}
            else:
                a = round(random.uniform(2, 10), 2)
                b = round(random.uniform(0.1, 5), 2)
                return {'question': f"{a} - {b} = ", 'answer': round(a - b, 2)}
        elif type_id == 'decimal_multiplication':
            if random.choice([True, False]):
                a = round(random.uniform(0.1, 10), 2)
                b = round(random.uniform(0.1, 1), 2)
                return {'question': f"{a} × {b} = ", 'answer': round(a * b, 2)}
            else:
                a = round(random.uniform(1, 10), 2)
                b = round(random.uniform(0.01, 0.2), 2)
                return {'question': f"{a} ÷ {b} = ", 'answer': round(a / b, 2)}
    elif category_id == 'fraction_operations':
        if type_id == 'fraction_addition':
            if random.choice([True, False]):
                a = random.randint(1, 9)
                b = random.randint(1, 9)
                c = random.randint(1, 9)
                return {'question': f"\\frac{{{a}}}{{{b}}} + \\frac{{{c}}}{{{b}}} = ", 'answer': f"{a+c}/{b}"}
            else:
                a = random.randint(1, 9)
                b = random.randint(2, 9)
                c = random.randint(1, 9)
                d = random.randint(2, 9)
                return {'question': f"\\frac{{{a}}}{{{b}}} + \\frac{{{c}}}{{{d}}} = ", 'answer': f"{a*d+c*b}/{b*d}"}
        elif type_id == 'fraction_multiplication':
            if random.choice([True, False]):
                a = random.randint(1, 9)
                b = random.randint(2, 9)
                c = random.randint(1, 9)
                d = random.randint(2, 9)
                return {'question': f"\\frac{{{a}}}{{{b}}} × \\frac{{{c}}}{{{d}}} = ", 'answer': f"{a*c}/{b*d}"}
            else:
                a = random.randint(1, 9)
                b = random.randint(2, 9)
                c = random.randint(1, 9)
                d = random.randint(2, 9)
                return {'question': f"\\frac{{{a}}}{{{b}}} ÷ \\frac{{{c}}}{{{d}}} = ", 'answer': f"{a*d}/{b*c}"}
        elif type_id == 'fraction_mixed':
            a = random.randint(1, 9)
            b = random.randint(2, 9)
            c = random.randint(1, 9)
            d = random.randint(2, 9)
            e = random.randint(1, 9)
            f = random.randint(2, 9)
            return {'question': f"\\frac{{{a}}}{{{b}}} × (\\frac{{{c}}}{{{d}}} + \\frac{{{e}}}{{{f}}}) = ", 'answer': f"{a*c*f+a*e*b*d}/{b*d*f}"}
    elif category_id == 'percentage_operations':
        if type_id == 'percentage_conversion':
            if random.choice([True, False]):
                a = random.choice([25, 50, 75, 100])
                return {'question': f"{a}% = ", 'answer': f"{a/100}"}
            else:
                a = random.choice([1, 2, 3, 4, 5, 10, 20, 25])
                return {'question': f"{a}/4 = ", 'answer': f"{a*25}%"}
        elif type_id == 'percentage_application':
            if random.choice([True, False]):
                a = random.randint(10, 100)
                b = random.choice([10, 20, 25, 50, 75])
                return {'question': f"{a}的{b}%是多少", 'answer': a * b / 100}
            else:
                a = random.randint(10, 100)
                b = random.randint(10, 100)
                c = random.randint(10, 100)
                return {'question': f"{a}比{b}多百分之几", 'answer': f"{(a-b)/b*100}%"}
    elif category_id == 'operation_laws':
        if type_id == 'simplified_calculation':
            if random.choice([True, False]):
                a = round(random.uniform(1, 10), 1)
                return {'question': f"{a} × 99 + {a} = ", 'answer': a * 100}
            else:
                a = random.randint(1, 9)
                b = random.randint(1, 9)
                c = random.randint(1, 9)
                return {'question': f"{a}/{b} × {c} + {a}/{b} × {b-c} = ", 'answer': f"{a*c+b*a*(b-c)}/{b}"}
    
    return {'question': '', 'answer': ''}

def generate_poetry_questions(poem_names, count=10):
    data_file = os.path.join('data', 'poetry.json')
    if not os.path.exists(data_file):
        return []
    
    with open(data_file, 'r', encoding='utf-8') as f:
        poetry_data = json.load(f)
    
    selected_poems = []
    for content in poetry_data.get('contents', []):
        if content.get('type') == '诗歌':
            for poem in content.get('works', []):
                if poem['name'] in poem_names:
                    selected_poems.append(poem)
    
    questions = []
    for poem in selected_poems:
        lines = poem['text']
        for i in range(0, len(lines) - 1, 2):
            if i + 1 < len(lines):
                if random.choice([True, False]):
                    question = {
                        'question': f"{lines[i]}，__________。",
                        'answer': lines[i + 1],
                        'poem_name': poem['name']
                    }
                else:
                    question = {
                        'question': f"__________，{lines[i + 1]}。",
                        'answer': lines[i],
                        'poem_name': poem['name']
                    }
                questions.append(question)
    
    random.shuffle(questions)
    return questions[:count]

def generate_english_questions(unit, direction, count=None):
    data_file = os.path.join('data', 'english.json')
    if not os.path.exists(data_file):
        return []
    
    with open(data_file, 'r', encoding='utf-8') as f:
        english_data = json.load(f)
    
    unit_data = None
    for u in english_data['units']:
        if u['name'] == unit:
            unit_data = u
            break
    
    if not unit_data:
        return []
    
    questions = []
    for word in unit_data['words']:
        if direction == 'cn_to_en':
            question = {
                'question': word['chinese'],
                'answer': word['english']
            }
        else:
            question = {
                'question': word['english'],
                'answer': word['chinese']
            }
        questions.append(question)
    
    return questions

def create_word_document(questions, title, cols=5, font_name='Times New Roman', chinese_font='宋体', font_size=14):
    doc = Document()
    
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.name = font_name
    title_run.font.size = Pt(16)
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
    title_run.font.bold = True
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=(len(questions) + cols - 1) // cols + 1, cols=cols)
    table.autofit = False
    table.style = 'Table Grid'
    
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top={"sz": "6", "val": "single", "color": "000000"},
                           bottom={"sz": "6", "val": "single", "color": "000000"},
                           left={"sz": "6", "val": "single", "color": "000000"},
                           right={"sz": "6", "val": "single", "color": "000000"})
    
    for i, q in enumerate(questions):
        row = i // cols + 1
        col = i % cols
        cell = table.cell(row, col)
        
        question_text = q['question']
        
        if '\\frac' in question_text:
            para = cell.paragraphs[0]
            
            import re
            frac_pattern = r'\\frac\{(\d+)\}\{(\d+)\}'
            matches = list(re.finditer(frac_pattern, question_text))
            
            last_pos = 0
            for match in matches:
                text_before = question_text[last_pos:match.start()]
                if text_before:
                    run = para.add_run(text_before)
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
                
                numerator = match.group(1)
                denominator = match.group(2)
                
                num_run = para.add_run(numerator)
                num_run.font.name = font_name
                num_run.font.size = Pt(font_size)
                num_run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
                
                slash_run = para.add_run('/')
                slash_run.font.name = font_name
                slash_run.font.size = Pt(font_size)
                slash_run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
                
                den_run = para.add_run(denominator)
                den_run.font.name = font_name
                den_run.font.size = Pt(font_size)
                den_run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
                
                last_pos = match.end()
            
            text_after = question_text[last_pos:]
            if text_after:
                run = para.add_run(text_after)
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
        else:
            para = cell.paragraphs[0]
            run = para.add_run(question_text)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
        
        para.paragraph_format.line_spacing = 1.5
    
    for cell in table.rows[0].cells:
        para = cell.paragraphs[0]
        run = para.add_run("")
        run.font.name = font_name
        run.font.size = Pt(font_size)
        para.paragraph_format.line_spacing = 1.5
    
    filename = f"output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join('data', filename)
    doc.save(filepath)
    
    return filepath

def create_english_word_document(questions, title, cols=3):
    doc = Document()
    
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(16)
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    title_run.font.bold = True
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=(len(questions) + cols - 1) // cols + 1, cols=cols)
    table.autofit = False
    table.style = 'Table Grid'
    
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top={"sz": "6", "val": "single", "color": "000000"},
                           bottom={"sz": "6", "val": "single", "color": "000000"},
                           left={"sz": "6", "val": "single", "color": "000000"},
                           right={"sz": "6", "val": "single", "color": "000000"})
    
    for i, q in enumerate(questions):
        row = i // cols + 1
        col = i % cols
        cell = table.cell(row, col)
        
        para = cell.paragraphs[0]
        run = para.add_run(q['question'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10.5)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        para.paragraph_format.line_spacing = 1.5
        
        answer_para = cell.add_paragraph()
        answer_run = answer_para.add_run("____________________")
        answer_run.font.name = 'Times New Roman'
        answer_run.font.size = Pt(10.5)
        answer_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        answer_para.paragraph_format.line_spacing = 1.5
    
    for cell in table.rows[0].cells:
        para = cell.paragraphs[0]
        run = para.add_run("")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10.5)
        para.paragraph_format.line_spacing = 1.5
    
    filename = f"output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join('data', filename)
    doc.save(filepath)
    
    return filepath

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/ai/chat')
def ai_chat_page():
    return render_template('ai_chat_new.html')

@app.route('/api/math/config')
def get_math_config():
    """获取数学配置，使用缓存"""
    cache_key = 'math_config'
    cached_data = get_cached_data(cache_key)
    if cached_data:
        return jsonify(cached_data)
    
    data_file = os.path.join('data', 'math_config.json')
    if os.path.exists(data_file):
        with open(data_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        set_cached_data(cache_key, data)
        return jsonify(data)
    return jsonify({})

@app.route('/api/poetry/list')
def get_poetry_list():
    """获取古诗列表，使用缓存"""
    cache_key = 'poetry_list'
    cached_data = get_cached_data(cache_key)
    if cached_data:
        return jsonify(cached_data)
    
    data_file = os.path.join('data', 'poetry.json')
    if os.path.exists(data_file):
        with open(data_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        poetry_list = []
        for content in data.get('contents', []):
            if content.get('type') == '诗歌':
                poetry_list.extend(content.get('works', []))
        set_cached_data(cache_key, poetry_list)
        return jsonify(poetry_list)
    return jsonify([])

@app.route('/api/english/units')
def get_english_units():
    """获取英语单元，使用缓存"""
    cache_key = 'english_units'
    cached_data = get_cached_data(cache_key)
    if cached_data:
        return jsonify(cached_data)
    
    data_file = os.path.join('data', 'english.json')
    if os.path.exists(data_file):
        with open(data_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        units_data = data.get('units', [])
        set_cached_data(cache_key, units_data)
        return jsonify(units_data)
    return jsonify([])

@app.route('/api/generate/math', methods=['POST'])
def generate_math():
    config = request.json
    questions = generate_math_questions(config)
    return jsonify({'questions': questions})

@app.route('/api/generate/poetry', methods=['POST'])
def generate_poetry():
    data = request.json
    poem_names = data.get('poems', [])
    count = data.get('count', 10)
    questions = generate_poetry_questions(poem_names, count)
    return jsonify({'questions': questions})

@app.route('/api/generate/english', methods=['POST'])
def generate_english():
    data = request.json
    unit = data.get('unit')
    direction = data.get('direction', 'cn_to_en')
    questions = generate_english_questions(unit, direction)
    return jsonify({'questions': questions})

@app.route('/api/download/math', methods=['POST'])
def download_math():
    config = request.json
    questions = generate_math_questions(config)
    filepath = create_word_document(questions, '数学练习题', cols=5, font_name='Times New Roman', chinese_font='宋体', font_size=14)
    return send_file(filepath, as_attachment=True)

@app.route('/api/download/poetry', methods=['POST'])
def download_poetry():
    data = request.json
    poem_names = data.get('poems', [])
    count = data.get('count', 10)
    questions = generate_poetry_questions(poem_names, count)
    filepath = create_word_document(questions, '古诗默写练习', cols=2, font_name='楷体', chinese_font='楷体')
    return send_file(filepath, as_attachment=True)

@app.route('/api/download/english', methods=['POST'])
def download_english():
    data = request.json
    unit = data.get('unit')
    direction = data.get('direction', 'cn_to_en')
    questions = generate_english_questions(unit, direction)
    filepath = create_english_word_document(questions, '英语默写练习', cols=3)
    return send_file(filepath, as_attachment=True)

@app.route('/api/download/report', methods=['POST'])
def download_report():
    data = request.json
    type_name = data.get('type', '练习')
    answers = data.get('answers', [])
    
    doc = Document()
    
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    title = doc.add_heading('练习报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    correct_count = sum(1 for a in answers if a.get('is_correct', False))
    total_count = len(answers)
    percentage = round((correct_count / total_count * 100), 1) if total_count > 0 else 0
    
    summary = doc.add_paragraph()
    summary.add_run(f'共 {total_count} 题，答对 {correct_count} 题，答错 {total_count - correct_count} 题，正确率 {percentage}%')
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    for i, answer in enumerate(answers, 1):
        p = doc.add_paragraph()
        p.add_run(f'第 {i} 题：').bold = True
        
        question_run = p.add_run(f'\n题目：{answer.get("question", "")}')
        question_run.font.name = '宋体'
        question_run.font.size = Pt(11)
        question_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        user_run = p.add_run(f'\n你的答案：{answer.get("user_answer", "")}')
        user_run.font.name = '宋体'
        user_run.font.size = Pt(11)
        user_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        correct_run = p.add_run(f'\n正确答案：{answer.get("correct_answer", "")}')
        correct_run.font.name = '宋体'
        correct_run.font.size = Pt(11)
        correct_run.font.bold = True
        correct_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        is_correct = answer.get('is_correct', False)
        result_run = p.add_run(f'\n结果：{"✓ 正确" if is_correct else "✗ 错误"}')
        result_run.font.name = '宋体'
        result_run.font.size = Pt(11)
        result_run.font.bold = True
        result_run.font.color.rgb = RGBColor(0, 128, 0) if is_correct else RGBColor(255, 0, 0)
        result_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(12)
    
    filename = f"练习报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join('data', filename)
    doc.save(filepath)
    
    return send_file(filepath, as_attachment=True)

ollama_process = None

def get_ollama_path():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    ollama_dir = os.path.join(script_dir, 'ollama')
    
    if platform.system().lower() == 'windows':
        return os.path.join(ollama_dir, 'ollama.exe')
    else:
        return os.path.join(ollama_dir, 'ollama')

def start_ollama():
    global ollama_process
    ollama_path = get_ollama_path()
    
    if os.path.exists(ollama_path):
        if ollama_process is None or ollama_process.poll() is not None:
            try:
                ollama_process = subprocess.Popen([ollama_path, 'serve'])
                return True
            except Exception as e:
                print(f"启动Ollama失败: {e}")
                return False
    else:
        print("未找到项目内Ollama，使用系统Ollama服务")
        return True

def stop_ollama():
    global ollama_process
    if ollama_process and ollama_process.poll() is None:
        ollama_process.terminate()
        ollama_process = None

@app.route('/api/ai/status', methods=['GET'])
def ai_status():
    try:
        models = ollama.list()
        return jsonify({'status': 'running', 'models': models})
    except Exception as e:
        return jsonify({'status': 'stopped', 'error': str(e)})

@app.route('/api/ai/chat', methods=['POST'])
def ai_chat():
    data = request.json
    question = data.get('question', '')
    model = data.get('model', 'qwen2.5:0.5b')
    
    if not question:
        return jsonify({'error': '请输入问题'}), 400
    
    def generate():
        try:
            print("使用系统Ollama服务")
            print(f"使用模型: {model}")
            
            response = ollama.chat(model=model, messages=[{
                'role': 'user',
                'content': question
            }], stream=True)
            
            for chunk in response:
                content = chunk.get('message', {}).get('content', '')
                if content:
                    # 立即yield，不做额外处理
                    yield f"data: {json.dumps({'content': content})}\n\n"
            
            yield "data: [DONE]\n\n"
            print("AI回答完成")
            
        except Exception as e:
            error_msg = str(e)
            print(f"AI回答失败: {error_msg}")
            if 'Failed to connect' in error_msg or 'Connection refused' in error_msg:
                yield f"data: {json.dumps({'content': '无法连接到Ollama服务。请确保Ollama已安装并正在运行。您可以运行 \"ollama serve\" 启动服务。详细说明请参考 OLLAMA_QUICKSTART.md'})}\n\n"
            else:
                yield f"data: {json.dumps({'content': f'回答失败：{error_msg}'})}\n\n"
            yield "data: [DONE]\n\n"
    
    return Response(stream_with_context(generate()), mimetype='text/event-stream', headers={
        'Cache-Control': 'no-cache',
        'X-Accel-Buffering': 'no'
    })

@app.route('/api/ai/models', methods=['GET'])
def ai_models():
    try:
        print("开始获取AI模型列表...")
        models = ollama.list()
        print(f"获取到的模型对象: {models}")
        print(f"模型对象类型: {type(models)}")
        
        models_list = []
        
        if hasattr(models, 'models') and models.models:
            print(f"找到 {len(models.models)} 个模型")
            for model in models.models:
                model_name = model.model if hasattr(model, 'model') else str(model)
                print(f"处理模型: {model_name}")
                models_list.append({
                    'name': model_name,
                    'size': model.size,
                    'modified': model.modified_at
                })
        else:
            print("未找到models属性或models为空")
        
        print(f"返回的模型列表: {models_list}")
        return jsonify({'models': models_list})
    except Exception as e:
        print(f"获取AI模型列表失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/ai/explain', methods=['POST'])
def ai_explain():
    data = request.json
    questions = data.get('questions', [])
    type_name = data.get('type', '')
    prompt = data.get('prompt', '')
    
    if not questions or len(questions) == 0:
        return jsonify({'error': '没有错题需要讲解'}), 400
    
    explanations = []
    
    try:
        for i, wrong_question in enumerate(questions):
            question_text = wrong_question.get('question', '')
            user_answer = wrong_question.get('userAnswer', '')
            correct_answer = wrong_question.get('correctAnswer', '')
            
            individual_prompt = f"""请详细讲解以下题目：

题目：{question_text}
学生答案：{user_answer or '未作答'}
正确答案：{correct_answer}

请给出详细的解题思路、知识点分析和正确答案的推导过程。如果学生的答案是错误的，请指出错误原因。"""
            
            try:
                print(f"正在讲解第 {i+1} 道错题...")
                response = ollama.chat(model='qwen2.5:0.5b', messages=[{
                    'role': 'user',
                    'content': individual_prompt
                }])
                
                explanation = response.get('message', {}).get('content', '')
                explanations.append(explanation)
                print(f"第 {i+1} 道错题讲解完成")
                
            except Exception as e:
                error_msg = str(e)
                print(f"第 {i+1} 道错题讲解失败: {error_msg}")
                if 'Failed to connect' in error_msg or 'Connection refused' in error_msg:
                    explanations.append(f"无法连接到Ollama服务。请确保Ollama已安装并正在运行。您可以运行 'ollama serve' 启动服务。详细说明请参考 OLLAMA_QUICKSTART.md")
                else:
                    explanations.append(f"讲解失败：{error_msg}")
        
        return jsonify({'explanations': explanations})
        
    except Exception as e:
        error_msg = str(e)
        if 'Failed to connect' in error_msg or 'Connection refused' in error_msg:
            return jsonify({'error': '无法连接到Ollama服务。请确保Ollama已安装并正在运行。您可以运行 "ollama serve" 启动服务。详细说明请参考 OLLAMA_QUICKSTART.md'}), 500
        return jsonify({'error': error_msg}), 500

@app.route('/api/ai/explain/stream', methods=['POST'])
def ai_explain_stream():
    data = request.json
    question_text = data.get('question', '')
    user_answer = data.get('userAnswer', '')
    correct_answer = data.get('correctAnswer', '')
    type_name = data.get('type', '')
    
    if not question_text:
        return jsonify({'error': '没有题目需要讲解'}), 400
    
    if type_name == 'math':
        subject_name = '数学'
    elif type_name == 'poetry':
        subject_name = '古诗'
    else:
        subject_name = '英语'
    
    prompt = f"""请详细讲解以下{subject_name}题目：

题目：{question_text}
学生答案：{user_answer or '未作答'}
正确答案：{correct_answer}

请给出详细的解题思路、知识点分析和正确答案的推导过程。如果学生的答案是错误的，请指出错误原因。"""
    
    def generate():
        try:
            print(f"开始流式讲解题目: {question_text[:50]}...")
            response = ollama.chat(model='qwen2.5:0.5b', messages=[{
                'role': 'user',
                'content': prompt
            }], stream=True)
            
            for chunk in response:
                content = chunk.get('message', {}).get('content', '')
                if content:
                    yield f"data: {json.dumps({'content': content})}\n\n"
            
            yield "data: [DONE]\n\n"
            print("流式讲解完成")
            
        except Exception as e:
            error_msg = str(e)
            print(f"流式讲解失败: {error_msg}")
            if 'Failed to connect' in error_msg or 'Connection refused' in error_msg:
                yield f"data: {json.dumps({'content': '无法连接到Ollama服务。请确保Ollama已安装并正在运行。'})}\n\n"
            else:
                yield f"data: {json.dumps({'content': f'讲解失败：{error_msg}'})}\n\n"
            yield "data: [DONE]\n\n"
    
    return Response(stream_with_context(generate()), mimetype='text/event-stream')

# ==================== 错题本功能 ====================

@app.route('/notebook')
def error_notebook_page():
    """错题本页面"""
    return render_template('error_notebook.html')

@app.route('/api/notebook/list', methods=['GET'])
def get_error_list():
    """获取错题列表"""
    data = get_error_notebook_data()
    subject = request.args.get('subject', '')
    
    errors = data.get('errors', [])
    
    # 按科目筛选
    if subject:
        errors = [e for e in errors if e.get('subject') == subject]
    
    # 按创建时间倒序排列
    errors.sort(key=lambda x: x.get('created_at', ''), reverse=True)
    
    return jsonify({'errors': errors, 'total': len(errors)})

@app.route('/api/notebook/subjects', methods=['GET'])
def get_subjects():
    """获取所有科目列表"""
    data = get_error_notebook_data()
    subjects = list(set(e.get('subject', '其他') for e in data.get('errors', [])))
    # 预设科目
    preset_subjects = ['数学', '语文', '英语', '物理', '化学', '生物', '历史', '地理', '政治', '其他']
    all_subjects = list(set(preset_subjects + subjects))
    all_subjects.sort()
    return jsonify({'subjects': all_subjects})

@app.route('/api/notebook/add', methods=['POST'])
def add_error():
    """添加错题"""
    try:
        # 获取表单数据
        title = request.form.get('title', '').strip()
        subject = request.form.get('subject', '其他').strip()
        question_text = request.form.get('question_text', '').strip()
        answer_text = request.form.get('answer_text', '').strip()
        notes = request.form.get('notes', '').strip()
        
        if not title:
            return jsonify({'error': '请输入错题标题'}), 400
        
        # 生成唯一ID
        error_id = str(uuid.uuid4())
        created_at = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        # 处理上传的文件
        question_files = []
        answer_files = []
        video_files = []
        
        # 处理错题文件（可多个）
        if 'question_files' in request.files:
            files = request.files.getlist('question_files')
            for file in files:
                if file and file.filename and allowed_file(file.filename):
                    # 生成安全的文件名
                    original_name = secure_filename(file.filename)
                    # 保留中文文件名，但添加唯一前缀
                    ext = original_name.rsplit('.', 1)[1].lower() if '.' in original_name else ''
                    new_filename = f"{error_id}_{uuid.uuid4().hex[:8]}.{ext}"
                    filepath = os.path.join(UPLOAD_FOLDER, 'questions', new_filename)
                    file.save(filepath)
                    question_files.append({
                        'filename': new_filename,
                        'original_name': file.filename,
                        'type': get_file_type(file.filename),
                        'size': os.path.getsize(filepath)
                    })
        
        # 处理答案文件（可多个）
        if 'answer_files' in request.files:
            files = request.files.getlist('answer_files')
            for file in files:
                if file and file.filename and allowed_file(file.filename):
                    original_name = secure_filename(file.filename)
                    ext = original_name.rsplit('.', 1)[1].lower() if '.' in original_name else ''
                    new_filename = f"{error_id}_{uuid.uuid4().hex[:8]}.{ext}"
                    filepath = os.path.join(UPLOAD_FOLDER, 'answers', new_filename)
                    file.save(filepath)
                    answer_files.append({
                        'filename': new_filename,
                        'original_name': file.filename,
                        'type': get_file_type(file.filename),
                        'size': os.path.getsize(filepath)
                    })
        
        # 处理视频文件（可多个）
        if 'video_files' in request.files:
            files = request.files.getlist('video_files')
            for file in files:
                if file and file.filename and allowed_file(file.filename):
                    original_name = secure_filename(file.filename)
                    ext = original_name.rsplit('.', 1)[1].lower() if '.' in original_name else ''
                    new_filename = f"{error_id}_{uuid.uuid4().hex[:8]}.{ext}"
                    filepath = os.path.join(UPLOAD_FOLDER, 'videos', new_filename)
                    file.save(filepath)
                    video_files.append({
                        'filename': new_filename,
                        'original_name': file.filename,
                        'type': get_file_type(file.filename),
                        'size': os.path.getsize(filepath)
                    })
        
        # 创建错题记录
        error_record = {
            'id': error_id,
            'title': title,
            'subject': subject,
            'question_text': question_text,
            'answer_text': answer_text,
            'notes': notes,
            'question_files': question_files,
            'answer_files': answer_files,
            'video_files': video_files,
            'created_at': created_at,
            'updated_at': created_at
        }
        
        # 保存到数据文件
        data = get_error_notebook_data()
        data['errors'].append(error_record)
        save_error_notebook_data(data)
        
        return jsonify({'success': True, 'id': error_id, 'message': '错题添加成功'})
        
    except Exception as e:
        print(f"添加错题失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'添加错题失败: {str(e)}'}), 500

@app.route('/api/notebook/get/<error_id>', methods=['GET'])
def get_error_detail(error_id):
    """获取错题详情"""
    data = get_error_notebook_data()
    for error in data.get('errors', []):
        if error.get('id') == error_id:
            return jsonify({'error': error})
    return jsonify({'error': '错题不存在'}), 404

@app.route('/api/notebook/update/<error_id>', methods=['POST'])
def update_error(error_id):
    """更新错题"""
    try:
        data = get_error_notebook_data()
        error_index = None
        
        for i, error in enumerate(data.get('errors', [])):
            if error.get('id') == error_id:
                error_index = i
                break
        
        if error_index is None:
            return jsonify({'error': '错题不存在'}), 404
        
        # 获取表单数据
        title = request.form.get('title', '').strip()
        subject = request.form.get('subject', '其他').strip()
        question_text = request.form.get('question_text', '').strip()
        answer_text = request.form.get('answer_text', '').strip()
        notes = request.form.get('notes', '').strip()
        
        if not title:
            return jsonify({'error': '请输入错题标题'}), 400
        
        # 更新基本信息
        data['errors'][error_index]['title'] = title
        data['errors'][error_index]['subject'] = subject
        data['errors'][error_index]['question_text'] = question_text
        data['errors'][error_index]['answer_text'] = answer_text
        data['errors'][error_index]['notes'] = notes
        data['errors'][error_index]['updated_at'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        # 处理新上传的文件
        if 'question_files' in request.files:
            files = request.files.getlist('question_files')
            for file in files:
                if file and file.filename and allowed_file(file.filename):
                    original_name = secure_filename(file.filename)
                    ext = original_name.rsplit('.', 1)[1].lower() if '.' in original_name else ''
                    new_filename = f"{error_id}_{uuid.uuid4().hex[:8]}.{ext}"
                    filepath = os.path.join(UPLOAD_FOLDER, 'questions', new_filename)
                    file.save(filepath)
                    data['errors'][error_index]['question_files'].append({
                        'filename': new_filename,
                        'original_name': file.filename,
                        'type': get_file_type(file.filename),
                        'size': os.path.getsize(filepath)
                    })
        
        if 'answer_files' in request.files:
            files = request.files.getlist('answer_files')
            for file in files:
                if file and file.filename and allowed_file(file.filename):
                    original_name = secure_filename(file.filename)
                    ext = original_name.rsplit('.', 1)[1].lower() if '.' in original_name else ''
                    new_filename = f"{error_id}_{uuid.uuid4().hex[:8]}.{ext}"
                    filepath = os.path.join(UPLOAD_FOLDER, 'answers', new_filename)
                    file.save(filepath)
                    data['errors'][error_index]['answer_files'].append({
                        'filename': new_filename,
                        'original_name': file.filename,
                        'type': get_file_type(file.filename),
                        'size': os.path.getsize(filepath)
                    })
        
        if 'video_files' in request.files:
            files = request.files.getlist('video_files')
            for file in files:
                if file and file.filename and allowed_file(file.filename):
                    original_name = secure_filename(file.filename)
                    ext = original_name.rsplit('.', 1)[1].lower() if '.' in original_name else ''
                    new_filename = f"{error_id}_{uuid.uuid4().hex[:8]}.{ext}"
                    filepath = os.path.join(UPLOAD_FOLDER, 'videos', new_filename)
                    file.save(filepath)
                    data['errors'][error_index]['video_files'].append({
                        'filename': new_filename,
                        'original_name': file.filename,
                        'type': get_file_type(file.filename),
                        'size': os.path.getsize(filepath)
                    })
        
        save_error_notebook_data(data)
        return jsonify({'success': True, 'message': '错题更新成功'})
        
    except Exception as e:
        print(f"更新错题失败: {str(e)}")
        return jsonify({'error': f'更新错题失败: {str(e)}'}), 500

@app.route('/api/notebook/delete/<error_id>', methods=['DELETE'])
def delete_error(error_id):
    """删除错题"""
    try:
        data = get_error_notebook_data()
        error_to_delete = None
        
        for error in data.get('errors', []):
            if error.get('id') == error_id:
                error_to_delete = error
                break
        
        if not error_to_delete:
            return jsonify({'error': '错题不存在'}), 404
        
        # 删除关联的文件
        for file_info in error_to_delete.get('question_files', []):
            filepath = os.path.join(UPLOAD_FOLDER, 'questions', file_info['filename'])
            if os.path.exists(filepath):
                os.remove(filepath)
        
        for file_info in error_to_delete.get('answer_files', []):
            filepath = os.path.join(UPLOAD_FOLDER, 'answers', file_info['filename'])
            if os.path.exists(filepath):
                os.remove(filepath)
        
        for file_info in error_to_delete.get('video_files', []):
            filepath = os.path.join(UPLOAD_FOLDER, 'videos', file_info['filename'])
            if os.path.exists(filepath):
                os.remove(filepath)
        
        # 从数据中删除
        data['errors'] = [e for e in data['errors'] if e.get('id') != error_id]
        save_error_notebook_data(data)
        
        return jsonify({'success': True, 'message': '错题删除成功'})
        
    except Exception as e:
        print(f"删除错题失败: {str(e)}")
        return jsonify({'error': f'删除错题失败: {str(e)}'}), 500

@app.route('/api/notebook/delete-file', methods=['POST'])
def delete_error_file():
    """删除错题的单个文件"""
    try:
        data_req = request.json
        error_id = data_req.get('error_id')
        file_type = data_req.get('file_type')  # 'question', 'answer', 'video'
        filename = data_req.get('filename')
        
        if not all([error_id, file_type, filename]):
            return jsonify({'error': '参数不完整'}), 400
        
        data = get_error_notebook_data()
        
        for error in data.get('errors', []):
            if error.get('id') == error_id:
                file_key = f'{file_type}_files'
                if file_key in error:
                    # 找到并删除文件记录
                    for i, file_info in enumerate(error[file_key]):
                        if file_info['filename'] == filename:
                            # 删除物理文件
                            folder_map = {'question': 'questions', 'answer': 'answers', 'video': 'videos'}
                            filepath = os.path.join(UPLOAD_FOLDER, folder_map[file_type], filename)
                            if os.path.exists(filepath):
                                os.remove(filepath)
                            # 删除记录
                            error[file_key].pop(i)
                            error['updated_at'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                            save_error_notebook_data(data)
                            return jsonify({'success': True, 'message': '文件删除成功'})
                
                return jsonify({'error': '文件不存在'}), 404
        
        return jsonify({'error': '错题不存在'}), 404
        
    except Exception as e:
        print(f"删除文件失败: {str(e)}")
        return jsonify({'error': f'删除文件失败: {str(e)}'}), 500

@app.route('/uploads/<folder>/<filename>')
def serve_upload(folder, filename):
    """提供上传文件的访问"""
    if folder not in ['questions', 'answers', 'videos']:
        return jsonify({'error': '无效的文件夹'}), 400
    
    filepath = os.path.join(UPLOAD_FOLDER, folder, filename)
    if os.path.exists(filepath):
        return send_file(filepath)
    return jsonify({'error': '文件不存在'}), 404

@app.route('/api/notebook/search', methods=['GET'])
def search_errors():
    """搜索错题"""
    keyword = request.args.get('keyword', '').strip()
    subject = request.args.get('subject', '')
    
    if not keyword and not subject:
        return get_error_list()
    
    data = get_error_notebook_data()
    errors = data.get('errors', [])
    
    results = []
    for error in errors:
        # 科目筛选
        if subject and error.get('subject') != subject:
            continue
        
        # 关键词搜索
        if keyword:
            searchable = f"{error.get('title', '')} {error.get('question_text', '')} {error.get('answer_text', '')} {error.get('notes', '')}"
            if keyword.lower() not in searchable.lower():
                continue
        
        results.append(error)
    
    results.sort(key=lambda x: x.get('created_at', ''), reverse=True)
    return jsonify({'errors': results, 'total': len(results)})

@app.route('/api/notebook/stats', methods=['GET'])
def get_notebook_stats():
    """获取错题本统计信息"""
    data = get_error_notebook_data()
    errors = data.get('errors', [])
    
    # 统计各科目错题数量
    subject_counts = {}
    for error in errors:
        subject = error.get('subject', '其他')
        subject_counts[subject] = subject_counts.get(subject, 0) + 1
    
    # 统计文件数量
    total_files = 0
    for error in errors:
        total_files += len(error.get('question_files', []))
        total_files += len(error.get('answer_files', []))
        total_files += len(error.get('video_files', []))
    
    return jsonify({
        'total_errors': len(errors),
        'subject_counts': subject_counts,
        'total_files': total_files
    })

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=3000)